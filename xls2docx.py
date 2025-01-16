import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import os
import time
import re
import pandas as pd
import xlrd
from openpyxl import load_workbook


# 需要 把所有修改点全部按格式修改

# 定义xls文件
default_path = 'E:/小说工作区/X岛——花园二号计划/workspace'  # 修改点1/4 默认工作路径，注意一定要把路径写全，分隔符用“/”
input_file_name = '55192991_网友面基，遇见气定大火球_20240625'  # 修改点2/4 替换为你的xls文件名字，不需要写扩展名
xls_file_path = default_path + '/' + f'{input_file_name}.xls'


# data = ILLEGAL_CHARACTERS_RE.sub(r'', data)  # 240625修改--清洗xls中的非法字符




# 将DataFrame保存为xlsx文件
df = pd.read_excel(xls_file_path)
xlsx_file_path = default_path + '/' + 'converted_file.xlsx'
# 修改点3/4 替换为你想要保存的xlsx文件路径，xlsx会作为中间文件。docx也会和它在同一个目录,一般不用改

df.to_excel(xlsx_file_path, index=False)

print("xlsx转换完成，xlsx文件已保存到", xlsx_file_path)


def xlsx_to_docx_with_images(xlsx_path, docx_path):
    # Load the Excel workbook
    wb = openpyxl.load_workbook(xlsx_path)
    ws = wb.active

    # Create a new Word document
    doc = Document()

    # Load the Excel file using pandas to find the "饼干" string
    df = pd.read_excel(xlsx_path)
    cookie_position = df.apply(lambda row: row.astype(str).str.contains('饼干').any(), axis=1)

    next_row_value = None
    if cookie_position.any():
        row_index = cookie_position[cookie_position].index[0]
        column_name = df.columns[df.iloc[row_index].astype(str).str.contains('饼干')][0]
        if row_index + 1 < len(df):
            next_row_value = df[column_name].iloc[row_index + 1]

    # Iterate through each row in the worksheet
    for row in ws.iter_rows(values_only=False):
        # Create a new paragraph for each row
        paragraph = doc.add_paragraph()

        for idx, cell in enumerate(row):
            if cell.value:
                run = paragraph.add_run(str(cell.value) + '\t')
                # Set font to 微软雅黑
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(12)  # Set default font size

                if idx == 1:  # Change font color for the first column (index 0)
                    run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green color
                elif idx == 2:
                    run = paragraph.add_run(':\n')
                elif idx == 3:  # Change font color for the fourth column (index 3)
                    run.font.color.rgb = RGBColor(139, 69, 19)  # Brownish red color

                # Check if the cell value matches the next_row_value
                if next_row_value and str(next_row_value) in str(cell.value):
                    po_run = paragraph.add_run('(PO主)\t')
                    po_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for "PO主"
                    po_run.font.name = '微软雅黑'
                    po_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    po_run.font.size = Pt(12)  # Set font size for "PO主"

        # Add the separator after each row
        doc.add_paragraph("%%====分割====%%")

    # Save the Word document
    doc.save(docx_path)


# Example usage
xlsx_path = xlsx_file_path
docx_path = default_path + '/' + f'{input_file_name}.docx'  # 修改点4/4 输出文件名，一般不用管
xlsx_to_docx_with_images(xlsx_path, docx_path)
print("docx转换完成，docx文件已保存到", docx_path)

# 延时2秒
time.sleep(1)
print("docx格式化完成，1秒后删除xlsx中间文件")

# 删除中间文件
os.remove(xlsx_path)
print("xlsx文件已删除")
