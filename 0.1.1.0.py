import urllib.request
from io import BytesIO
import gzip
from bs4 import BeautifulSoup
import re
import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox
from datetime import datetime
import os
import time
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import openpyxl
from openpyxl.drawing.image import Image as ExcelImage
from tqdm import tqdm

# 正则表达式对象
findthreadid = re.compile(r'<li><a href="/Admin/Content/sagePost/id/(.*?).html">SAGE</a></li>')
findthreaduid = re.compile(r'<span class="h-threads-info-uid">ID:(.*?)</span>')
findcreatedat = re.compile(r'<span class="h-threads-info-createdat">(.*?)</span>')
findcontent = re.compile(r'<div class="h-threads-content">\n(.*?)</div>', re.S)
findmaintitle = re.compile(r'<h2 class="h-title">(.*?)</h2>')
findmainemail = re.compile(r'<span class="h-threads-info-email">(.*?)</span>')

def main():
    root = tk.Tk()
    root.withdraw()
    
    # 获取用户Cookie
    cookie = get_user_cookie()
    if not cookie:
        print("未提供Cookie，程序终止。")
        return
    
    # 选择包含串号的TXT文件
    file_path = filedialog.askopenfilename(title="选择包含串号的TXT文件", filetypes=[("Text files", "*.txt")])
    if not file_path:
        print("未选择文件，程序终止。")
        return
    
    # 读取串号列表
    with open(file_path, 'r', encoding='utf-8') as f:
        thread_ids = [line.strip() for line in f.readlines() if line.strip()]
    
    # 遍历处理每个串号
    for thread_id in thread_ids:
        print(f"\n开始处理串号：{thread_id}")
        try:
            process_thread(thread_id, cookie)
        except Exception as e:
            print(f"处理串号 {thread_id} 时发生错误：{str(e)}")

def get_user_cookie():
    """获取用户输入的Cookie，优先使用缓存"""
    CACHE_DIR = os.path.join(os.getenv('APPDATA'), 'CookieCache')
    COOKIE_PATH = os.path.join(CACHE_DIR, 'cookie.txt')
    os.makedirs(CACHE_DIR, exist_ok=True)

    old_cookie = None
    if os.path.exists(COOKIE_PATH):
        with open(COOKIE_PATH, 'r', encoding='utf-8') as f:
            old_cookie = f.read().strip()

    use_new = True
    if old_cookie:
        use_new = messagebox.askyesno("Cookie设置", "是否使用新Cookie？")

    if use_new:
        cookie = simpledialog.askstring("输入Cookie", "请输入Cookie值：")
        if cookie is None:
            messagebox.showinfo("提示", "未输入Cookie，程序终止。")
            return None
        while not cookie.strip():
            messagebox.showwarning("输入错误", "Cookie不能为空，请重新输入。")
            cookie = simpledialog.askstring("输入Cookie", "请输入Cookie值：")
            if cookie is None:
                messagebox.showinfo("提示", "未输入Cookie，程序终止。")
                return None
        with open(COOKIE_PATH, 'w', encoding='utf-8') as f:
            f.write(cookie.strip())
    else:
        cookie = old_cookie

    return cookie

def process_thread(thread_id, cookie):
    baseurl = f'https://www.nmbxd1.com/t/{thread_id}?page='
    datalist, author_id, main_title, main_email = get_data(baseurl, cookie)
    save_date = datetime.now().strftime('%Y%m%d')
    
    # 创建输出目录
    output_dir = f"{thread_id}_{save_date}"
    os.makedirs(output_dir, exist_ok=True)
    
    # 保存Excel文件（带图片）
    xlsx_path = os.path.join(output_dir, f"{main_title}_{save_date}.xlsx")
    save_data_to_xlsx(datalist, xlsx_path, main_title, main_email)
    
    # 生成DOCX文档
    docx_path = os.path.join(output_dir, f"{main_title}_{save_date}.docx")
    generate_docx(datalist, docx_path, thread_id, author_id, save_date)
    
    # 下载图片并插入Excel
    download_and_insert_images(datalist, output_dir, cookie, xlsx_path)

def ask_url(url, cookie):
    head = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:126.0) Gecko/20100101 Firefox/126.1",
        "Cookie": cookie
    }
    req = urllib.request.Request(url, headers=head)
    try:
        with urllib.request.urlopen(req) as response:
            h = response.read()
            buff = BytesIO(h)
            f = gzip.GzipFile(fileobj=buff)
            return f.read().decode('utf-8')
    except Exception as e:
        print(f"请求URL时发生错误：{str(e)}")
        return ""

def parse_page(html):
    soup = BeautifulSoup(html, "html.parser")
    main_item = soup.find('div', class_="h-threads-item-main")
    items = soup.find_all('div', class_="h-threads-item-reply-main")
    theme = soup.find('h2', class_="h-title")
    return main_item, items, theme

def parse_item(item):
    data = []
    item_str = str(item)
    
    # 提取基础信息
    threadid = re.findall(findthreadid, item_str)
    data.append(threadid[0] if threadid else "")
    
    threaduid = re.findall(findthreaduid, item_str)
    data.append(threaduid[0] if threaduid else "")
    
    createdat = re.findall(findcreatedat, item_str)
    data.append(createdat[0] if createdat else "")
    
    # 处理内容
    content = re.findall(findcontent, item_str)
    if content:
        content = content[0]
        content = re.sub(r'<.*?>', '', content)
        content = re.sub(r'&gt;', "＞", content)
        content = re.sub(r'&lt;', "＜", content)
        content = re.sub(r'<font color="#789922">', "", content)
        content = re.sub(r'</font>', "", content)
        data.append(content.strip())
    else:
        data.append("")
    
    # 提取图片URL
    soup = BeautifulSoup(item_str, 'html.parser')
    img_urls = [img['src'] for img in soup.find_all('img', class_='h-threads-img') if img.get('src')]
    data.append('; '.join(img_urls))
    
    return data

def get_data(baseurl, cookie):
    datalist = []
    page = 1
    author_id = main_title = main_email = ""
    
    while True:
        url = baseurl + str(page)
        html = ask_url(url, cookie)
        if not html:
            break
            
        main_item, items, theme = parse_page(html)
        
        if page == 1 and main_item:
            data = parse_item(main_item)
            datalist.append(data)
            author_id = data[1]
            
            # 提取主标题和邮箱
            main_title_match = re.findall(findmaintitle, str(theme))
            main_title = main_title_match[0] if main_title_match else ""
            main_email_match = re.findall(findmainemail, str(main_item))
            main_email = main_email_match[0] if main_email_match else ""
        
        if page > 1 and len(items) <= 1:
            break
            
        for item in items:
            data = parse_item(item)
            if data[0] != "9999999":
                datalist.append(data)
                
        page += 1
        time.sleep(1)  # 控制爬取速度
        
    return datalist, author_id, main_title, main_email

def save_data_to_xlsx(datalist, save_path, main_title, main_email):
    df = pd.DataFrame(datalist, columns=["串号", "饼干", "时间", "内容", "图片URL"])
    
    # 使用上下文管理器确保正确保存
    with pd.ExcelWriter(save_path, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='主数据')
        
        # 获取工作簿和工作表对象
        workbook = writer.book
        worksheet = writer.sheets['主数据']
        
        # 添加标题信息
        worksheet.cell(row=1, column=1, value="标题")
        worksheet.cell(row=1, column=2, value=main_title)
        worksheet.cell(row=2, column=1, value="邮箱")
        worksheet.cell(row=2, column=2, value=main_email)
        
        # 设置列宽
        worksheet.column_dimensions['A'].width = 15
        worksheet.column_dimensions['B'].width = 20
        worksheet.column_dimensions['C'].width = 25
        worksheet.column_dimensions['D'].width = 60
        worksheet.column_dimensions['E'].width = 40

    print(f"Excel文件已保存至：{save_path}")

def generate_docx(datalist, docx_path, thread_id, author_id, save_date):
    doc = Document()
    
    # 添加标题信息
    doc.add_heading(f'串号：{thread_id}', level=1)
    doc.add_paragraph(f'作者ID：{author_id}')
    doc.add_paragraph(f'保存时间：{save_date}')
    doc.add_page_break()
    
    # 添加内容
    for data in datalist:
        p = doc.add_paragraph()
        p.add_run(f"【{data[2]}】").bold = True
        p.add_run(f" ID:{data[1]}：\n")
        p.add_run(data[3]).font.color.rgb = RGBColor(0x78, 0x99, 0x22)  # X岛特色绿色
        
        if data[1] == author_id:  # 修正后的正确行
            po_run = p.add_run(" (PO主)")
            po_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            po_run.bold = True
        
        doc.add_paragraph("%%====分割线====%%", style='Intense Quote')
    
    # 设置中文字体
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    doc.save(docx_path)

def download_and_insert_images(datalist, output_dir, cookie, xlsx_path):
    # 创建图片目录
    img_dir = os.path.join(output_dir, "images")
    os.makedirs(img_dir, exist_ok=True)
    
    # 下载图片并记录路径
    img_paths = []
    for data in tqdm(datalist, desc="下载图片"):
        img_urls = data[4].split('; ')
        for idx, url in enumerate(img_urls):
            if not url:
                continue
                
            try:
                filename = f"{data[0]}_{idx}.jpg"
                save_path = os.path.join(img_dir, filename)
                
                req = urllib.request.Request(url, headers={"Cookie": cookie})
                with urllib.request.urlopen(req) as response:
                    with open(save_path, 'wb') as f:
                        f.write(response.read())
                        
                img_paths.append({
                    "post_id": data[0],
                    "img_path": save_path
                })
            except Exception as e:
                print(f"下载失败：{url}，错误：{e}")
    
    # 插入图片到Excel
    wb = openpyxl.load_workbook(xlsx_path)
    ws = wb.active
    
    # 添加图片列
    ws.insert_cols(5)
    ws.cell(row=1, column=5, value="图片")
    
    for row_idx, data in enumerate(datalist, start=2):
        post_id = data[0]
        images = [img['img_path'] for img in img_paths if img['post_id'] == post_id]
        
        if images:
            cell = ws.cell(row=row_idx, column=5)
            cell.value = "\n".join(images)
            
            # 插入图片（示例插入第一张）
            try:
                img = ExcelImage(images[0])
                img.width = 100
                img.height = 100
                ws.add_image(img, f"E{row_idx}")
            except Exception as e:
                print(f"插入图片失败：{e}")
    
    wb.save(xlsx_path)

if __name__ == "__main__":
    main()
    print("所有串号处理完毕！")