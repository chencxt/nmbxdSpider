###
# 这是一个用于爬取x岛的串的python爬虫
# 本文件由爬虫0.0.1.6c.py和xls2docx.py合并而来。时间2024年06月25日
# 你需要修改的部分一共有6处：cookie和第二部分的四个修改点。其中，有4个修改点不需要改动。第五个修改点位于126行附近
# 请时刻确保自己的网络畅通，关闭梯子。否则有概率报错
# 1.1新增功能 xls非法字符过滤；自定义标题
# 20240712 修改输出文件格式 使其适合列表分割
# 20240703 添加调用 img_spider.py 用于生成含有图像的xlsx文件。默认文件路径为脚本所在路径，暂时不改。必须确保工作文件夹内有此脚本文件和img_spider.py
# img_spider.py的修改点也需要更改，folder2xlsx.py的文件路径也需要更改
###
import urllib.request
from io import BytesIO
import gzip
from bs4 import BeautifulSoup
import re
import xlwt
import tkinter as tk
from tkinter import simpledialog
from datetime import datetime
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import pandas as pd
import os
import time
from bs4 import BeautifulSoup
import openpyxl
from openpyxl.drawing.image import Image
import re
import os
import time
from tqdm import tqdm

savepath = ""  # 全局变量，用于保存路径
thread_id = ""
main_title = ""
save_date = ""

# 创建正则表达式对象，表示规范（字符串的模式）
findthreadid = re.compile('<li><a href="/Admin/Content/sagePost/id/(.*?).html">SAGE</a></li>')
findthreaduid = re.compile('<span class="h-threads-info-uid">ID:(.*?)</span>')
findcreatedat = re.compile('<span class="h-threads-info-createdat">(.*?)</span>')
findcontent = re.compile('<div class="h-threads-content">\n(.*?)</div>', re.S)
findmaintitle = re.compile('<h2 class="h-title">(.*?)</h2>')
findmainemail = re.compile('<span class="h-threads-info-email">(.*?)</span>')


def main():
    root = tk.Tk()
    root.withdraw()
    thread_id = simpledialog.askstring("输入", "请输入需要爬取的串号:").strip()
    themeuseful = simpledialog.askstring("输入", "请输入想要优先覆盖的标题，空置则默认使用扒到的内容:").strip()
    if not thread_id:
        print("未输入串号，程序终止。")
        return

    print('开始爬取......')
    baseurl = f'https://www.nmbxd1.com/t/{thread_id}?page='
    datalist, author_id, main_title, main_email = getData(baseurl, themeuseful)

    save_date = datetime.now().strftime('%Y%m%d')
    savepath = f"{main_title} - {save_date}.xls"
    txt_filename = f"{main_title} - {save_date}.txt"
    saveData(datalist, savepath, thread_id, author_id, save_date, main_title, main_email, txt_filename)
    xls2docx(thread_id, main_title, save_date)


def askURL(url):
    head = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:126.0) Gecko/20100101 Firefox/126.1",
        "Cookie": "PHPSESSID=qdaqmuds779s9o8nmhutv0ghl1; memberUserspapapa=%E0an%0C%B5%03%BE%E67%B6%88%02iq%E0%14c%D7_%E57%96Y%D4%FD%04%5E%EB%E5%A6%E6%D5%91%9D%1F%5E%99w%BF%F9K%B4L%88%2F%5BV%AF%C64%DB%5D%89%24%11%B8%93%01%9A%22%C5%D8%93%19; userhash=%E7%DAe%DE%DB%DAV%AA%3D%7DKL%EF%28%E0%16G%CA%1F%B9%40%99%5B%E3"
    }

    req = urllib.request.Request(url, headers=head)
    html = ""
    with urllib.request.urlopen(req) as response:
        h = response.read()
        buff = BytesIO(h)
        f = gzip.GzipFile(fileobj=buff)
        html = f.read().decode('utf-8')
    return html


def parsePage(html):
    soup = BeautifulSoup(html, "html.parser")
    theme = soup.find('h2', class_="h-title")
    main_item = soup.find('div', class_="h-threads-item-main")
    items = soup.find_all('div', class_="h-threads-item-reply-main")
    return main_item, items, theme


def parseItem(item):
    ILLEGAL_CHARACTERS_RE = re.compile(r'[\000-\010]|[\013-\014]|[\016-\037]')
    data = []
    item = str(item)
    threadid = re.findall(findthreadid, item)
    threadid = threadid[0] if threadid else ""
    data.append(threadid)

    threaduid = re.findall(findthreaduid, item)
    data.append(threaduid[0] if threaduid else "")

    createdat = re.findall(findcreatedat, item)
    data.append(createdat[0] if createdat else "")

    content = re.findall(findcontent, item)
    if content:
        content = content[0]
        content = re.sub("<.*?>", "", content)
        content = re.sub('&gt;', "＞", content)
        content = re.sub('<font color="#789922">&gt;&gt;', ">>", content)
        content = re.sub('&lt;', "＜", content)
        content = re.sub("<br/>", "", content)
        content = re.sub("<br>", "", content)
        content = re.sub("</br>", "", content)
        content = re.sub('<font color="#789922">', "", content)
        content = re.sub("</font>", "", content)
        content = re.sub("<b>", "", content)
        content = re.sub("</b>", "", content)
        content = re.sub("</small>", "", content)
        content = re.sub("<small>", "", content)
        content = ILLEGAL_CHARACTERS_RE.sub(r'', content)  # 240625修改--清洗非法字符
        data.append(content.strip())
    else:
        data.append("")

    return data


def getData(baseurl, themeuseful):
    datalist = []
    page = 1
    author_id = None
    main_title = ""
    main_email = ""
    while True:
        url = baseurl + str(page)
        html = askURL(url)
        print("第%d页" % page)
        time.sleep(1)  # 修改点5 控制爬虫速度，一般不用改
        main_item, items, theme = parsePage(html)

        if page == 1 and main_item:
            data = parseItem(main_item)
            main_title = re.findall(findmaintitle, str(theme))[0] if re.findall(findmaintitle,
                                                                                str(theme)) else ""
            # themeuseful = simpledialog.askstring("输入", "请输入想要优先覆盖的标题，空置则默认使用扒到的内容:").strip()
            if themeuseful:
                # 定义要替换的字符串
                original_string = main_title

                # 定义正则表达式模式，匹配所有数字
                pattern = r'- (.*?) -'

                # 定义替换后的字符串，这里将所有数字替换为'num'
                replacement = '- ' + themeuseful + ' -'

                # 使用re.sub()进行替换
                main_title = re.sub(pattern, replacement, original_string)

            main_email = re.findall(findmainemail, str(main_item))[0] if re.findall(findmainemail,
                                                                                    str(main_item)) else ""
            datalist.append(data)
            author_id = data[1]

        if page > 1 and len(items) <= 1:
            break
        for item in items:
            data = parseItem(item)
            if data[0] != "9999999":
                datalist.append(data)
        page += 1
    return datalist, author_id, main_title, main_email


def contains_chinese(text):
    for ch in text:
        if '\u4e00' <= ch <= '\u9fff':
            return True
    return False


def saveData(datalist, savepath, thread_id, author_id, save_date, main_title, main_email, txt_filename):
    print("正在保存数据......")
    workbook = xlwt.Workbook(encoding="utf-8", style_compression=0)
    worksheet = workbook.add_sheet('Sheet1', cell_overwrite_ok=True)

    style_english = xlwt.XFStyle()
    font_english = xlwt.Font()
    font_english.name = 'Times New Roman'
    font_english.bold = False
    style_english.font = font_english

    style_chinese = xlwt.XFStyle()
    font_chinese = xlwt.Font()
    font_chinese.name = '宋体'
    font_chinese.bold = False
    style_chinese.font = font_chinese

    col = ("串号", "饼干", "时间", "内容")
    for i in range(0, len(col)):
        worksheet.write(2, i, col[i], style_chinese)

    worksheet.write(0, 0, "标题", style_chinese)
    worksheet.write(1, 0, main_title, style_chinese)
    worksheet.write(0, 1, "副标题", style_chinese)
    worksheet.write(1, 1, main_email, style_chinese)

    for i in range(0, len(datalist)):
        print(f"正在写入第{i + 1}条数据")
        data = datalist[i]
        for j in range(0, len(data)):
            if contains_chinese(data[j]):
                worksheet.write(i + 3, j, data[j], style_chinese)
            else:
                worksheet.write(i + 3, j, data[j], style_english)

    workbook.save(savepath)
    print("数据保存成功！")

    txt_savepath = os.path.join(os.getcwd(), txt_filename)
    with open(txt_savepath, 'w', encoding='utf-8') as f:
        f.write(f"串号：{thread_id}\n")
        f.write(f"po：{author_id}\n")
        f.write(f"保存时间：{save_date}\n")
        f.write("\n======\n\n")
        f.write(f"{main_title}\n")
        f.write(f"{main_email}\n\n")

        for data in datalist:
            if data[1] == author_id:
                f.write(f"{data[3]}\n\n\n\n")
    print(f"作者信息保存成功！路径: {txt_savepath}")
    # 延时1秒
    time.sleep(1)
    print("爬虫爬取完成，1秒后进入格式化文档部分")


'''
从这里开始，进入文件转换部分————————————————————————————————————————————————————————————


    需要 把所有修改点全部按格式修改
    '''


def xls2docx(thread_id, main_title, save_date):
    # 读取xls文件
    default_path = 'E:/小说工作区/X岛——花园二号计划/workspace'  # 修改点1/4 默认工作路径，注意一定要把路径写全，分隔符用“/”
    input_file_name = f'{main_title} - {save_date}'  # 修改点2/4 替换为你的xls文件名字，不需要写扩展名，不用改
    xls_file_path = default_path + '/' + f"{input_file_name}.xls"
    df = pd.read_excel(xls_file_path)

    # 将DataFrame保存为xlsx文件
    xlsx_file_path = default_path + '/' + 'converted_file.xlsx'
    # 修改点3/4 替换为你想要保存的xlsx文件路径，xlsx会作为中间文件。docx也会和它在同一个目录,不用改

    df.to_excel(xlsx_file_path, index=False)

    print("xlsx转换完成，xlsx文件已保存到", xlsx_file_path)

    def xlsx_to_docx_with_images(xlsx_path, docx_path):
        # Load the Excel workbook
        wb = openpyxl.load_workbook(xlsx_path)
        ws = wb.active

        # Create a new Word document
        doc = Document()

        # Load the Excel file using pandas to find the "饼干" string
        df = pd.read_excel(xlsx_path)
        cookie_position = df.apply(lambda row: row.astype(str).str.contains('饼干').any(), axis=1)

        next_row_value = None
        if cookie_position.any():
            row_index = cookie_position[cookie_position].index[0]
            column_name = df.columns[df.iloc[row_index].astype(str).str.contains('饼干')][0]
            if row_index + 1 < len(df):
                next_row_value = df[column_name].iloc[row_index + 1]

        # Iterate through each row in the worksheet
        for row in ws.iter_rows(values_only=False):
            # Create a new paragraph for each row
            paragraph = doc.add_paragraph()

            for idx, cell in enumerate(row):
                if cell.value:
                    run = paragraph.add_run(str(cell.value) + '\t')
                    # Set font to 微软雅黑
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(12)  # Set default font size

                    if idx == 1:  # Change font color for the first column (index 0)
                        run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green color
                    elif idx == 2:
                        run = paragraph.add_run(':\n')
                    elif idx == 3:  # Change font color for the fourth column (index 3)
                        run.font.color.rgb = RGBColor(139, 69, 19)  # Brownish red color

                    # Check if the cell value matches the next_row_value
                    if next_row_value and str(next_row_value) in str(cell.value):
                        po_run = paragraph.add_run('(PO主)\t')
                        po_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for "PO主"
                        po_run.font.name = '微软雅黑'
                        po_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                        po_run.font.size = Pt(12)  # Set font size for "PO主"

            # Add the separator after each row
            doc.add_paragraph("%%====分割====%%")

        # Save the Word document
        doc.save(docx_path)

    # Example usage
    xlsx_path = xlsx_file_path
    docx_path = default_path + '/' + f"{input_file_name}.docx"  # 修改点4/4 输出文件名，不用管
    xlsx_to_docx_with_images(xlsx_path, docx_path)
    print("docx转换完成，docx文件已保存到", docx_path)

    # 延时1秒
    time.sleep(1)
    print("docx格式化完成，1秒后删除xlsx中间文件")

    # 删除中间文件
    os.remove(xlsx_path)
    print("xlsx文件已删除")


if __name__ == "__main__":
    main()
    print("主爬虫执行完毕！进入图像拔取环节")
    SleepTime = 3
    while SleepTime >= 1:
        time.sleep(1)
        print(f"{SleepTime}秒后开始进入图像拔取环节")
        SleepTime -= 1
    # main.py
    with open('img_spider.py', encoding='utf-8') as f:
        code = f.read()

    exec(code)
    print("程序全部执行完成\nxls是原始文字扒取文件 - - txt是只含po主的文本文件 - - docx是加工后的文本文件 - - xlsx是图像打包文件")

